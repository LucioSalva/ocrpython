"""DOCX export."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.shared import Pt


def export_docx(
    *,
    text_content: str | None,
    extracted_fields: dict[str, Any] | None,
    document_info: dict[str, Any],
    output_path: Path,
) -> Path:
    doc = DocxDocument()

    title = doc.add_heading("Resultado OCR", level=1)
    title.alignment = 1  # center

    info = doc.add_paragraph()
    info.add_run("Archivo original: ").bold = True
    info.add_run(str(document_info.get("original_filename", "")))

    info2 = doc.add_paragraph()
    info2.add_run("Plantilla: ").bold = True
    info2.add_run(str(document_info.get("template_code") or "—"))

    info3 = doc.add_paragraph()
    info3.add_run("Idioma detectado: ").bold = True
    info3.add_run(str(document_info.get("language") or "—"))

    if extracted_fields:
        doc.add_heading("Campos extraídos", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Campo"
        hdr[1].text = "Valor"
        for k, v in extracted_fields.items():
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = "" if v is None else str(v)

    doc.add_heading("Texto extraído", level=2)
    body_para = doc.add_paragraph()
    run = body_para.add_run(text_content or "")
    run.font.size = Pt(10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
